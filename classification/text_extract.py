"""
classification/text_extract.py

Part 2 - Tier 2 helper: extract plain text from a downloaded primary-data file
so the rule-based classifier can read its content.

Supported:
    plain text  .txt .csv .tsv .tab .md .json .xml .html .htm
    code/syntax  .r .do .sps .sas .py .yml .yaml        (treated as text)
    .rtf         via striprtf
    .pdf         via pypdf
    .docx        via python-docx
    .zip / .qdpx recursively extract text from the text-capable files inside
                 (QDA project files are zip containers of primary data)

Binary formats with no usable text (.zsav .sav .dta .rds .rdata .xls .xlsx
.jpg .png ...) return "" -> the caller leaves that file unclassified, which is
the honest outcome (there is no readable primary-data text).

Everything is wrapped in try/except so a single bad file never breaks a batch.
Extracted text is capped so classification stays fast on huge files.
"""

from __future__ import annotations

import io
import zipfile
import pathlib

MAX_CHARS = 200_000          # cap text handed to the classifier
_ZIP_MEMBER_LIMIT = 40       # don't read more than N members from one archive

TEXT_EXTS = {
    ".txt", ".csv", ".tsv", ".tab", ".md", ".json", ".xml",
    ".html", ".htm", ".r", ".do", ".sps", ".sas", ".py",
    ".yml", ".yaml", ".log", ".dat",
}
ZIP_EXTS = {".zip", ".qdpx", ".qdp", ".qde"}
BINARY_EXTS = {
    ".zsav", ".sav", ".dta", ".rds", ".rdata", ".rda", ".xls", ".xlsx",
    ".jpg", ".jpeg", ".png", ".gif", ".tiff", ".bmp", ".mp3", ".mp4",
    ".wav", ".avi", ".mov", ".bin", ".sas7bdat", ".por",
}


def _clip(text: str) -> str:
    return text[:MAX_CHARS]


def _read_textbytes(data: bytes) -> str:
    for enc in ("utf-8", "latin-1"):
        try:
            return data.decode(enc, errors="ignore")
        except Exception:
            continue
    return ""


def _from_pdf(path: pathlib.Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    out = []
    for page in reader.pages:
        try:
            out.append(page.extract_text() or "")
        except Exception:
            continue
        if sum(len(x) for x in out) > MAX_CHARS:
            break
    return "\n".join(out)


def _from_docx(path: pathlib.Path) -> str:
    import docx
    d = docx.Document(str(path))
    return "\n".join(p.text for p in d.paragraphs)


def _from_rtf(path: pathlib.Path) -> str:
    from striprtf.striprtf import rtf_to_text
    return rtf_to_text(path.read_text(encoding="latin-1", errors="ignore"))


def _from_zip(path: pathlib.Path) -> str:
    out = []
    try:
        with zipfile.ZipFile(path) as zf:
            names = [n for n in zf.namelist() if not n.endswith("/")]
            for name in names[:_ZIP_MEMBER_LIMIT]:
                ext = pathlib.Path(name).suffix.lower()
                try:
                    with zf.open(name) as fh:
                        data = fh.read()
                except Exception:
                    continue
                if ext in TEXT_EXTS or ext in {".rtf", ""}:
                    out.append(_read_textbytes(data))
                elif ext == ".xml":
                    out.append(_read_textbytes(data))
                # nested pdf/docx inside zips: write to a temp buffer
                elif ext == ".pdf":
                    try:
                        from pypdf import PdfReader
                        r = PdfReader(io.BytesIO(data))
                        out.append("\n".join((p.extract_text() or "") for p in r.pages))
                    except Exception:
                        pass
                elif ext == ".docx":
                    try:
                        import docx
                        d = docx.Document(io.BytesIO(data))
                        out.append("\n".join(p.text for p in d.paragraphs))
                    except Exception:
                        pass
                if sum(len(x) for x in out) > MAX_CHARS:
                    break
    except (zipfile.BadZipFile, Exception):
        return ""
    return "\n".join(out)


def extract_text(path) -> str:
    """Return best-effort plain text for a file, or '' if none is available."""
    path = pathlib.Path(path)
    if not path.exists() or not path.is_file():
        return ""
    ext = path.suffix.lower()

    try:
        if ext in BINARY_EXTS:
            return ""
        if ext == ".pdf":
            return _clip(_from_pdf(path))
        if ext == ".docx":
            return _clip(_from_docx(path))
        if ext == ".rtf":
            return _clip(_from_rtf(path))
        if ext in ZIP_EXTS:
            return _clip(_from_zip(path))
        if ext in TEXT_EXTS:
            return _clip(_read_textbytes(path.read_bytes()))
        # unknown extension: try as text, give up quietly if it looks binary
        raw = path.read_bytes()[:MAX_CHARS]
        if b"\x00" in raw[:1024]:
            return ""
        return _read_textbytes(raw)
    except Exception:
        return ""


if __name__ == "__main__":
    import sys
    t = extract_text(sys.argv[1])
    print(f"[{len(t)} chars]\n{t[:800]}")
